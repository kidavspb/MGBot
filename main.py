import telebot
import datetime
import time
from docx import Document
import subprocess
from pdf2image import convert_from_path
from PIL import Image
import time
import os
import traceback
import logging
from secret_file import *

logging.basicConfig(
    level=logging.INFO,
    #    filename = "MGlog.log",
    format="%(asctime)s - %(module)s - %(levelname)s - %(funcName)s: %(lineno)d - %(message)s, user: %(user)s",
    handlers=[
        logging.FileHandler("MGlog.log"),
        logging.StreamHandler()
    ],
    datefmt='%H:%M:%S',
)

while True:
    try:
        logging.info("Bot running..", extra={"user": 0})
        months = ["января", "февраля",
                  "марта", "апреля", "мая",
                  "июня", "июля", "августа",
                  "сентября", "октября", "ноября",
                  "декабря"]
        bot = telebot.TeleBot(TOKEN)


        @bot.message_handler(commands=['start'])
        def welcome_message(message):
            logging.info("Start message from user", extra={"user": message.from_user.username})
            bot.send_message(message.chat.id, "Здравствуйте! Пожалуйста, выберите тип справки:\n\n"
                                              "<b>Электронная</b> — подходит для большинства случаев и делается практически мгновенно\n\n"
                                              "<b>Бумажная</b> — делается долго, потому что на нее ставится живая подпись с печатью "
                                              "(заказывайте только если уверены, что нужна именно бумажная)",
                             parse_mode='HTML', reply_markup=create_response_markup_type())


        @bot.message_handler(commands=['cancel'])
        def cancel_message(message):
            logging.info("Cancel message from user", extra={"user": message.from_user.username})
            bot.send_message(message.chat.id, "Отменено")
            bot.register_next_step_handler(message, welcome_message)


        def create_response_markup_type():
            markup = telebot.types.InlineKeyboardMarkup()
            markup.add(telebot.types.InlineKeyboardButton("💻 Электронная", callback_data="Digital"),
                       telebot.types.InlineKeyboardButton("📃 Бумажная", callback_data="Paper"))
            return markup


        def get_full_name(message, isDigital):
            full_name = message.text
            bot.send_message(message.chat.id,
                             "Хорошо! Теперь, пожалуйста, введите причину для справки в формате «принимает участие в ...»")
            bot.register_next_step_handler(message, get_reason, full_name=full_name, isDigital=isDigital)


        def get_reason(message, full_name, isDigital):
            start_length = len("принимает участие в")
            try:
                if len(message.text) <= start_length:
                    raise Exception("подозрительно короткая причина")
                reason = message.text[0].lower() + message.text[1:]
                if reason[:start_length] != "принимает участие в":
                    raise Exception("причина должна начинаться со слов «принимает участие в ...»")

            except Exception as e:
                bot.send_message(message.chat.id, f"Не подходит: {e}. Попробуйте еще раз")
                bot.register_next_step_handler(message, get_reason, full_name=full_name, isDigital=isDigital)
            else:
                bot.send_message(message.chat.id, "Почти готово! Осталось уточнить, когда было это мероприятие")
                bot.register_next_step_handler(message, get_date, full_name=full_name, reason=reason,
                                               isDigital=isDigital)


        def get_date(message, full_name, reason, isDigital):
            words = "".join(c for c in message.text.replace(".", " ") if c.isalnum() or c == " ").strip().split()
            year = current_year = int(datetime.date.today().strftime("%Y"))

            try:
                day = int(words[0])
                if words[1].isdigit():
                    month = months[int(words[1]) - 1]
                elif words[1] in months:
                    month = words[1]
                else:
                    raise Exception("Месяц странный")

                if len(words) > 2:
                    year = int(words[2])
                    if year < 100:
                        year += current_year // 100 * 100
                    if abs(year - int(datetime.date.today().strftime("%Y"))) > 2:
                        raise Exception("Год странный")
            except Exception as e:
                bot.send_message(message.chat.id, f"Что-то не то с датой, попробуйте еще раз\n\n<code>{e}</code>",
                                 parse_mode='HTML')
                bot.register_next_step_handler(message, get_date, full_name=full_name, reason=reason,
                                               isDigital=isDigital)
            else:
                date = f"{day} {month} {year} года"
                bot.send_message(message.chat.id,
                                 "Спасибо! Справка отправлена на согласование, нужно немного подождать")

                logging.info("Send request", extra={"user": message.from_user.username})
                user_username = "@" + message.from_user.username
                user_name = message.from_user.first_name + " " + message.from_user.last_name
                info_text = f"📨 Заявка на справку от {user_username} ({user_name})"
                user_message = f"<b>Тип:</b> {'электронная' if isDigital else 'бумажная'}\n\n<b>ФИО:</b> {full_name}\n\n<b>Причина:</b> {reason}\n\n<b>Дата:</b> {date} "

                # Send message to bot administrator
                bot.send_message(ADMIN_CHAT_ID,
                                 f"{message.chat.id} {int(time.time())}\n\n{info_text}\n\n{user_message}",
                                 parse_mode='HTML', reply_markup=create_response_markup_approval())


        def create_response_markup_approval():
            markup = telebot.types.InlineKeyboardMarkup()
            markup.add(telebot.types.InlineKeyboardButton("✅ Одобрить", callback_data="Approve"),
                       telebot.types.InlineKeyboardButton("❌ Отклонить", callback_data="Reject"))
            return markup


        @bot.callback_query_handler(func=lambda call: True)
        def callback(call):
            if call.data == "Digital":
                bot.send_message(call.message.chat.id,
                                 "Хороший выбор! Пожалуйста, введите теперь полное ФИО человека, которому нужна справка")
                bot.register_next_step_handler(call.message, get_full_name, isDigital=True)
            elif call.data == "Paper":
                bot.send_message(call.message.chat.id,
                                 "Хорошо! Пожалуйста, введите теперь полное ФИО человека, которому нужна справка")
                bot.register_next_step_handler(call.message, get_full_name, isDigital=False)

            elif call.data == "Reject":
                bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      text=call.message.text + "\n\n❌ Отклонено")
                bot.send_message(call.message.text.split('\n')[0].split(' ')[0],
                                 "К сожалению, справка не была согласована 🤷‍")
                logging.info("Reject request", extra={"user": call.message.text.split('\n')[0].split(' ')[0]})
            elif call.data == "Approve":
                # bot.send_message(CHANNEL_ID, call.message.text)
                bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      text=call.message.text + "\n\n✅ Одобрено")
                bot.send_message(call.message.text.split('\n')[0].split(' ')[0], "Ура, справка согласована 🎉")
                logging.info("Approve request", extra={"user": call.message.text.split('\n')[0].split(' ')[0]})

                full_name = (call.message.text.split('ФИО: '))[1].split('\n')[0]
                reason = (call.message.text.split('Причина: '))[1].split('\n')[0]
                date = (call.message.text.split('Дата: '))[1].split('\n')[0]

                isDigital = True if (call.message.text.split('Тип: '))[1].split('\n')[0] == "электронная" else False

                def create_PDF_certificate(full_name, reason, date, isDigital):
                    # Open the Word document
                    if isDigital:
                        document = Document("original_file.docx")
                    else:
                        document = Document("original_file_paper.docx")

                    today = datetime.date.today().strftime("%d.%m.%Y")
                    # Modify the document
                    document.paragraphs[3].text = f"От «{today.split('.')[0]}» {months[int(today.split('.')[1]) - 1]} {today.split('.')[2]} г."
                    document.paragraphs[10].text = document.paragraphs[10].text.replace('Иванов Иван Иванович',  full_name)
                    document.paragraphs[10].text = document.paragraphs[10].text.replace('1 декабря 2022 года', date)
                    document.paragraphs[10].text = document.paragraphs[10].text.replace('приминает участие в', reason)

                    # Save the changes to the document
                    document.save("edited_file.docx")

                    # Convert .docx to .pdf
                    def generate_pdf(doc_path, path):
                        subprocess.call(['soffice',
                                         # '--headless',
                                         '--convert-to',
                                         'pdf',
                                         '--outdir',
                                         path,
                                         doc_path])
                        return doc_path

                    generate_pdf("edited_file.docx", "")

                create_PDF_certificate(full_name, reason, date, isDigital)

                if isDigital:
                    # Convert .pdf to .png
                    pages = convert_from_path('edited_file.pdf', 500)
                    pages[0].save('edited_file.png', 'PNG')

                    result_filename = full_name + ".pdf"

                    # Convert .png to .pdf
                    img = Image.open("edited_file.png")
                    img.save(result_filename)

                    # Send modified file to the user
                    with open(result_filename, "rb") as f:
                        seconds = int(time.time() - int(call.message.text.split('\n')[0].split(' ')[1]))
                        h = seconds // 3600
                        m = seconds % 3600 // 60
                        s = seconds % 3600 % 60
                        if h > 0:
                            preparation_time = '{} ч. {} мин. {} сек.'.format(h, m, s)
                        elif m > 0:
                            preparation_time = '{} мин. {} сек.'.format(m, s)
                        elif s > 0:
                            preparation_time = '{} сек.'.format(s)

                        # time.sleep(10) # wait for the file to be saved

                        id_to_user = bot.send_document(CHANNEL_ID, f, caption=f"Сделано за {preparation_time}").document.file_id
                        bot.send_document(call.message.text.split('\n')[0].split(' ')[0], id_to_user)
                        logging.info("Send to user",
                                     extra={"user": call.message.text.split('\n')[0].split(' ')[0]})
                else:
                    bot.send_message(call.message.text.split('\n')[0],
                                     "Уже печатаем справку. Сообщим, как только она будет готова 👌")
                    # result_filename = "НА ПЕЧАТЬ - " + full_name + ".pdf"
                    result_filename = "edited_file.pdf"
                    with open(result_filename, "rb") as f:
                        bot.send_document(PRINTER_CHAT_ID, f,
                                          caption=call.message.text.split('\n')[0]
                                                  + f"\n\nПечать справки для {full_name}",
                                          reply_markup=telebot.types.InlineKeyboardMarkup().add(
                                              telebot.types.InlineKeyboardButton("🖨️ Напечатано",
                                                                                 callback_data="Printed")))

                # Сlean up after ourselves
                os.remove(result_filename)

            elif call.data == "Printed":
                bot.edit_message_caption(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                         caption=call.message.caption + "\n\n🖨️ Напечатано")
                bot.send_message(call.message.caption.split('\n')[0], f"Cправка для " +
                                 (call.message.caption.split('Печать справки для '))[1].split('\n')[0] +
                                 " готова, можно забирать в " + PLACE)
                logging.info("Send to user", extra={"user": call.message.caption.split('\n')[0]})


        logging.info("Bot running..", extra={"user": 0})
        bot.polling(none_stop=True)

    except Exception as e:
        logging.error(e, extra={"user": 0})
        bot.stop_polling()

        time.sleep(15)

        logging.info("Running again!", extra={"user": 0})

# bot.polling()
