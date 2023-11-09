import json
from docx import Document
import datetime
import subprocess

months = ["января", "февраля",
          "марта", "апреля", "мая",
          "июня", "июля", "августа",
          "сентября", "октября", "ноября",
          "декабря"]

path_to_docs_file = "docs/"


def create_PDF_certificate(full_name, reason, date, isDigital):
    # Open the Word document
    if isDigital:
        document = Document(path_to_docs_file+"original_file.docx")
    else:
        document = Document(path_to_docs_file+"original_file_paper.docx")

    today = datetime.date.today().strftime("%d.%m.%Y")
    # Modify the document
    document.paragraphs[3].text = f"От «{today.split('.')[0]}» {months[int(today.split('.')[1]) - 1]} {today.split('.')[2]} г."
    document.paragraphs[10].text = document.paragraphs[10].text.replace('Иванов Иван Иванович',
                                                                        full_name)
    document.paragraphs[10].text = document.paragraphs[10].text.replace('1 декабря 2022 года', date)
    document.paragraphs[10].text = document.paragraphs[10].text.replace('приминает участие в', reason)

    # Save the changes to the document
    document.save(path_to_docs_file+"edited_file.docx")

    generate_pdf(path_to_docs_file+"edited_file.docx", path_to_docs_file)


# Convert .docx to .pdf
def generate_pdf(doc_path, path):
    subprocess.call(['soffice',
                     # '--headless',
                     '--convert-to',
                     'pdf',
                     '--outdir',
                     path,
                     doc_path])
    return doc_path


def count_average_time(isDigital, current_time=0):
    path_to_data_file = "data/"
    type = "digital" if isDigital else "paper"

    with open(path_to_data_file+"time.json", 'r') as fr:
        data = json.load(fr)

        average_time = data[type]["average_time"]
        number_of_requests = data[type]["number_of_requests"]

        if current_time == 0:
            return average_time
        else:
            data[type]["average_time"] = (average_time * number_of_requests + current_time) / (number_of_requests + 1)
            data[type]["number_of_requests"] = number_of_requests + 1

            with open(path_to_data_file+"time.json", "w") as fw:
                json.dump(data, fw, indent=2, ensure_ascii=False)


def beautiful_time(seconds):
    seconds = int(seconds)
    d = seconds // (3600 * 24)
    h = seconds % (3600 * 24) // 3600
    m = seconds % 3600 // 60
    s = seconds % 60

    preparation_time = ''

    if d > 0:
        preparation_time = f'{d} д.'
    if h > 0:
        if preparation_time: preparation_time += ' '
        preparation_time += f'{h} ч.'
    if m > 0:
        if preparation_time: preparation_time += ' '
        preparation_time += f'{m} мин.'
    if s > 0:
        if preparation_time: preparation_time += ' '
        preparation_time += f'{s} сек.'

    return preparation_time
