from docx import Document
import datetime
import subprocess

months = ["января", "февраля",
          "марта", "апреля", "мая",
          "июня", "июля", "августа",
          "сентября", "октября", "ноября",
          "декабря"]


def create_PDF_certificate(full_name, reason, date, isDigital):
    # Open the Word document
    if isDigital:
        document = Document("original_file.docx")
    else:
        document = Document("original_file_paper.docx")

    today = datetime.date.today().strftime("%d.%m.%Y")
    # Modify the document
    document.paragraphs[
        3].text = f"От «{today.split('.')[0]}» {months[int(today.split('.')[1]) - 1]} {today.split('.')[2]} г."
    document.paragraphs[10].text = document.paragraphs[10].text.replace('Иванов Иван Иванович',
                                                                        full_name)
    document.paragraphs[10].text = document.paragraphs[10].text.replace('1 декабря 2022 года', date)
    document.paragraphs[10].text = document.paragraphs[10].text.replace('приминает участие в', reason)

    # Save the changes to the document
    document.save("edited_file.docx")

    generate_pdf("edited_file.docx", "")


# Convert .docx to .pdf
def generate_pdf(doc_path, path):
    subprocess.call(['soffice',
                     # '--headless',
                     '--convert-to',
                     'pdf',
                     '--outdir',
                     path,
                     doc_path])
    return doc_path


def count_average_time(isDigital, current_time=0):
    if isDigital:
        filename = "digital_time.txt"
    else:
        filename = "paper_time.txt"

    with open(filename, "r") as tf:
        tf_content = tf.read().split('\n')
        avegage_time = int(float(tf_content[0]))
        number_of_requests = int(tf_content[1])

    if current_time == 0:
        return avegage_time

    avegage_time = (avegage_time * number_of_requests + current_time) / (number_of_requests + 1)
    number_of_requests += 1
    with open(filename, "w") as tf:
        tf.write(str(avegage_time) + '\n' + str(number_of_requests))


def beautiful_time(seconds):
    seconds = int(seconds)
    d = seconds // (3600 * 24)
    h = seconds % (3600 * 24) // 3600
    m = seconds % 3600 // 60
    s = seconds % 60

    preparation_time = ''

    if d > 0:
        preparation_time = f'{d} д.'
    if h > 0:
        if preparation_time: preparation_time += ' '
        preparation_time += f'{h} ч.'
    if m > 0:
        if preparation_time: preparation_time += ' '
        preparation_time += f'{m} мин.'
    if s > 0:
        if preparation_time: preparation_time += ' '
        preparation_time += f'{s} сек.'

    return preparation_time
